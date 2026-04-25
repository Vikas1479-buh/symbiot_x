import re
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import cv2
from docx import Document

# PATHS
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\Users\vikas\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin'

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# CASCADE
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
eye_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_eye.xml')


# 🔍 DETECTION
def detect_sensitive_data(text):
    text = text.replace("\n", " ")

    patterns = {
        "Aadhaar": r'\b\d{4}\s?\d{4}\s?\d{4}\b',
        "PAN": r'\b[A-Z]{5}[0-9]{4}[A-Z]\b',
        "Phone": r'\b[6-9]\d{9}\b',
        "Email": r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b',
        "Credit Card": r'\b(?:\d{4}[-\s]?){3}\d{4}\b',
        "Passport": r'\b[A-Z][0-9]{7}\b',
        "IFSC": r'\b[A-Z]{4}0[A-Z0-9]{6}\b',
        "Bank Account": r'\b\d{9,18}\b'
    }

    detected = {}
    for key, pattern in patterns.items():
        matches = re.findall(pattern, text)
        if matches:
            detected[key] = list(set(matches))

    return detected


# 🔥 MRZ DETECTION
def detect_mrz(text):
    return re.findall(r'\b[A-Z0-9<]{15,}\b', text)


# 🖼 MASK IMAGE
def mask_image(filepath, values, mask_eyes=False):
    img = cv2.imread(filepath)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    # 🔒 TEXT MASK
    chunks = []
    for v in values:
        clean = v.replace(" ", "")
        if len(clean) > 12:
            chunks.append(clean)
        else:
            chunks += [clean[i:i+4] for i in range(0, len(clean), 4)]

    for i in range(len(data['text'])):
        word = data['text'][i]
        for chunk in chunks:
            if chunk and chunk in word:
                x = data['left'][i]
                y = data['top'][i]
                w = data['width'][i]
                h = data['height'][i]
                cv2.rectangle(img, (x,y), (x+w,y+h), (0,0,0), -1)

    # 🔥 QR BLUR
    edges = cv2.Canny(gray, 100, 200)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    for cnt in contours:
        x,y,w,h = cv2.boundingRect(cnt)
        aspect_ratio = w / float(h)
        area = w * h

        if (0.9 < aspect_ratio < 1.1 and 80 < w < 250 and area > 5000):
            roi = img[y:y+h, x:x+w]
            edge_density = cv2.countNonZero(
                cv2.Canny(roi, 50, 150)
            ) / (w * h)

            if edge_density > 0.25:
                img[y:y+h, x:x+w] = cv2.GaussianBlur(roi, (51,51), 30)

    # 👁️ EYE MASK (ROBUST)
    if mask_eyes:
        faces = face_cascade.detectMultiScale(gray, 1.1, 4)

        for (x,y,w,h) in faces:
            roi_color = img[y:y+h, x:x+w]
            roi_gray = gray[y:y+h, x:x+w]

            eyes = eye_cascade.detectMultiScale(roi_gray, 1.1, 4)

            if len(eyes) >= 2:
                for (ex,ey,ew,eh) in eyes[:2]:
                    roi_color[ey:ey+eh, ex:ex+ew] = cv2.GaussianBlur(
                        roi_color[ey:ey+eh, ex:ex+ew], (35,35), 30
                    )
            else:
                # fallback eye band
                y1 = int(h * 0.25)
                y2 = int(h * 0.45)
                roi_color[y1:y2, :] = cv2.GaussianBlur(
                    roi_color[y1:y2, :], (51,51), 30
                )

    out = "masked_" + os.path.basename(filepath)
    cv2.imwrite(os.path.join(UPLOAD_FOLDER, out), img)
    return out


# ROUTES
@app.route('/image/<filename>')
def image(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)


@app.route('/download/<filename>')
def download(filename):
    path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(path):
        return jsonify({"error": "File not found"}), 404
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)


@app.route('/upload', methods=['POST'])
def upload():
    try:
        file = request.files['file']
        file_type = request.form.get("type")
        mask_eyes = request.form.get("mask_eyes") == "true"

        path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(path)

        # 🖼 IMAGE
        if file_type == "image":
            text = pytesseract.image_to_string(Image.open(path))
            detected = detect_sensitive_data(text)

            values = []
            for v in detected.values():
                values.extend(v)

            values.extend(detect_mrz(text))

            masked = mask_image(path, values, mask_eyes)

            return jsonify({
                "preview": f"http://127.0.0.1:5000/image/{masked}",
                "download": f"http://127.0.0.1:5000/download/{masked}"
            })

        # 📄 PDF
        elif file_type == "pdf":
            images = convert_from_path(path, poppler_path=POPPLER_PATH)

            masked_imgs = []
            for i, img in enumerate(images):
                temp = os.path.join(UPLOAD_FOLDER, f"temp_{i}.png")
                img.save(temp)

                text = pytesseract.image_to_string(img)
                detected = detect_sensitive_data(text)

                values = []
                for v in detected.values():
                    values.extend(v)

                values.extend(detect_mrz(text))

                masked_name = mask_image(temp, values, mask_eyes)
                masked_imgs.append(Image.open(os.path.join(UPLOAD_FOLDER, masked_name)))

            pdf_name = "masked_output.pdf"
            pdf_path = os.path.join(UPLOAD_FOLDER, pdf_name)

            masked_imgs[0].save(pdf_path, save_all=True, append_images=masked_imgs[1:])

            return jsonify({
                "download": f"http://127.0.0.1:5000/download/{pdf_name}"
            })

        # 📝 TXT (FIXED)
        elif file_type == "txt":
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()

            detected = detect_sensitive_data(text)

            for vals in detected.values():
                for val in vals:
                    text = text.replace(val, "XXXX")

            filename = "masked_" + file.filename
            out_path = os.path.join(UPLOAD_FOLDER, filename)

            with open(out_path, 'w', encoding='utf-8') as f:
                f.write(text)

            return jsonify({
                "download": f"http://127.0.0.1:5000/download/{filename}"
            })

        # 📘 DOCX (FIXED)
        elif file_type == "docx":
            doc = Document(path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            detected = detect_sensitive_data(full_text)

            for vals in detected.values():
                for val in vals:
                    full_text = full_text.replace(val, "XXXX")

            filename = "masked_" + file.filename
            out_path = os.path.join(UPLOAD_FOLDER, filename)

            new_doc = Document()
            new_doc.add_paragraph(full_text)
            new_doc.save(out_path)

            return jsonify({
                "download": f"http://127.0.0.1:5000/download/{filename}"
            })

        return jsonify({"error": "Invalid type"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)